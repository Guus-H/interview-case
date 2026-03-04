"""
Run this script once to generate the meeting transcript DOCX file.

Usage:
    pip install python-docx
    python setup/generate_docx.py
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "..", "transcripts", "meeting_transcript.docx")

MEETING_META = {
    "title": "Meeting Transcript — Meridian Employee Portal Requirements",
    "date": "4 March 2026",
    "duration": "~13 minutes",
    "platform": "Microsoft Teams",
}

PARTICIPANTS = [
    ("Sarah Chen", "Project Manager", "Plat4mation"),
    ("Mark de Vries", "Technical Consultant", "Plat4mation"),
    ("Lisa Janssen", "Functional Consultant", "Plat4mation"),
    ("Tom Bakker", "IT Manager", "Meridian Group"),
    ("Priya Sharma", "Brand & Communications Manager", "Meridian Group"),
]

TRANSCRIPT = [
    ("Sarah Chen", "Okay I think we're all in — yes, Priya just joined. Welcome everyone."),
    ("Priya Sharma", "Sorry, Teams was being difficult again."),
    ("Tom Bakker", "No worries. Can everyone hear me okay?"),
    ("Mark de Vries", "Loud and clear, Tom."),
    ("Sarah Chen", "Great. So this is our second requirements session for the Meridian employee portal project. Last time we covered the high-level scope — today I want to get into the specifics of what we're actually building. Tom, you said you had some things you wanted to kick off with?"),
    ("Tom Bakker", "Yeah thanks Sarah. So my main concern — and I've been hearing this from employees since we announced the project — is that whatever people see when they first log in actually makes sense. Right now they're jumping between three different intranet pages and two separate ticketing systems. So I want the landing page to be the one place where everything starts. You land there and you immediately see what's relevant to you."),
    ("Lisa Janssen", "That makes sense. Are you thinking more of a dashboard feel, or a classic intranet homepage?"),
    ("Tom Bakker", "Dashboard, definitely. I want people to land there and see — okay so there are a few things. Company news and announcements, because right now HR sends an email and half the people miss it. And then also some kind of quick overview of their open requests and pending tasks. So they're not having to go hunting for things."),
    ("Priya Sharma", "I love that direction. From a communications standpoint, having announcements front and center is huge for us. We've been asking for a proper internal comms channel for two years. And can I just add — it should be visual. Not just a bullet list of text."),
    ("Tom Bakker", "Exactly, yes. Cards with a featured banner at the top would be ideal."),
    ("Mark de Vries", "That's very doable in the Service Portal. We can set up a homepage widget that pulls news items from a dedicated announcements table. I'd go with a dedicated table rather than the knowledge base — gives you more control over scheduling and audience targeting."),
    ("Sarah Chen", "Good. So homepage: dashboard layout, announcements widget, summary of open requests. Let's keep that captured. Mark, you have that?"),
    ("Mark de Vries", "Already noted."),
    ("Lisa Janssen", "Before we go further on the homepage — I want to make sure we also get the service catalog on the table today, because I think that's where we have the most complexity. Tom, you mentioned two ticketing systems. One for IT and one for HR I assume? And the idea is to consolidate that into a single request experience?"),
    ("Tom Bakker", "Correct. Right now if you need a new laptop you go to one system. If you need to update your personal details or request leave, that's a different HR portal. Employees find it confusing. We want one front door. You come in, you see categories — IT, HR, Facilities — pick what you need and fill in a form."),
    ("Lisa Janssen", "And the forms — how complex are we talking? Are there fields that change depending on the category? Mandatory attachments, that kind of thing?"),
    ("Tom Bakker", "For some of them yes. If you're requesting new hardware you'd specify the type, and for anything above a certain value you'd need a manager approval before it moves forward. But I don't want to overcomplicate the UI. The form should feel simple even if there's logic behind it."),
    ("Mark de Vries", "That's standard catalog item behavior — dynamic fields, mandatory conditions, approval flows triggered by thresholds. We've done that before. The main decision is really about how you structure the categories and which fields are global versus per-item."),
    ("Priya Sharma", "Can I jump in on the design side for a second? I want to make sure the request form feels on-brand. We have a style guide, and our primary action color is a specific teal — not the default ServiceNow blue."),
    ("Mark de Vries", "Sure, we can theme the portal. That's CSS configuration, not a big lift."),
    ("Priya Sharma", "Good. And — this might sound like a small thing — where does the submit button go on the form?"),
    ("Tom Bakker", "At the bottom, right? Like every other form on the internet."),
    ("Priya Sharma", "Well, our brand guidelines actually say primary action buttons should be top-right. It's about reducing eye travel. The button is always in the same place, so users know where to look regardless of how long the form is."),
    ("Tom Bakker", "That feels really unintuitive to me. You haven't even filled in the form yet and the submit button is already visible? What if you scroll? Are there two buttons?"),
    ("Priya Sharma", "It would be sticky — it stays fixed in the top-right as you scroll. It's actually a pretty modern pattern. A lot of SaaS products work this way now."),
    ("Mark de Vries", "I've seen both. Sticky top-right can work for short forms, but for longer forms it feels disconnected from the content. What about a sticky footer bar instead? Always visible, but anchored at the bottom where users expect an action."),
    ("Priya Sharma", "Hmm. I could live with that as long as it uses the brand color and spans the full width of the form area."),
    ("Tom Bakker", "That works for me. Sarah, should we capture that?"),
    ("Sarah Chen", "Let's note it as a UI preference for the design phase — sticky submit button, full-width footer style, primary brand color. Not a functional requirement as such. Can we get back to the catalog functionality?"),
    ("Lisa Janssen", "Yes — one thing I want to confirm Tom: do employees need to be able to attach files when submitting a request? I'm thinking of the hardware scenario — you might want to attach a quote or an existing approval email."),
    ("Tom Bakker", "Definitely. File attachments are a must. And there should be a free-text notes field too — so people can give context even when there isn't a specific form field for their situation."),
    ("Lisa Janssen", "Perfect. So for the request catalog: category browsing, dynamic form fields per request type, mandatory fields, file attachment support, free-text notes. I think that covers the core of that one."),
    ("Sarah Chen", "Great, that's well scoped."),
    ("Mark de Vries", "While we're on the topic — I want to raise something related but slightly different. One of the patterns we see a lot is that employees submit tickets for things that are already answered in existing documentation. We get five tickets a week just about VPN connectivity. If we have a searchable knowledge base that's surfaced prominently — and ideally integrated with the request form — you can deflect a real chunk of that volume. Tom, does Meridian have existing knowledge content or would this be starting from scratch?"),
    ("Tom Bakker", "We have some stuff in Confluence — pretty outdated. And there's a SharePoint site that nobody uses because nothing is findable. So yeah, effectively greenfield. But I love the idea. If employees can find answers themselves before submitting a ticket, that's a big win for the IT team."),
    ("Mark de Vries", "ServiceNow has knowledge management built in. You get categories, a search interface, and you can even surface suggested articles when someone starts filling in a request — like, you're asking about VPN access, and the form shows three articles that might answer it first."),
    ("Priya Sharma", "Oh that's really smart. And from a comms angle, we could use knowledge articles for onboarding materials, company policies — it doesn't have to be just IT content."),
    ("Tom Bakker", "Good point. HR has been wanting a place to publish policy documents that people can actually find. As long as we can have different sections with different owners — IT manages IT content, HR manages HR content — I'm very much in favour."),
    ("Mark de Vries", "Exactly. Each knowledge category can have its own owner group with publish rights. And employees can rate whether an article was helpful — which is useful for keeping the content quality up over time."),
    ("Sarah Chen", "Good. So knowledge base — searchable, categorized, category-level ownership, helpfulness rating, and integrated article suggestions during request submission. Another requirement."),
    ("Lisa Janssen", "Should the article suggestions during submission be explicitly in scope, or is that a nice-to-have?"),
    ("Tom Bakker", "In scope. If it deflects even twenty percent of unnecessary tickets it pays for itself. I want that."),
    ("Sarah Chen", "Noted, in scope."),
    ("Priya Sharma", "Quick thing on the homepage — I didn't hear us mention quick links. Like shortcuts to the most commonly used things."),
    ("Tom Bakker", "Yes, good catch. Our office manager was very specific about this. She wants a kind of top-ten list — the things 80% of employees need 80% of the time. New request, check my open tickets, find a policy, maybe book a meeting room."),
    ("Lisa Janssen", "Should those quick links be configurable per role? An IT admin might want different shortcuts than a regular employee."),
    ("Tom Bakker", "Ideally yes, but keep it simple. A default set that works for everyone, and then a separate set for IT admins. Two tiers is enough for now — I don't want to turn this into a configuration nightmare."),
    ("Sarah Chen", "Okay so homepage gets updated — add a quick links widget, two role tiers. Got it."),
    ("Mark de Vries", "There's one more thing I think we need to cover and that's the approvals flow. We said some requests trigger a manager approval. Where does the manager actually go to handle that? Right now they get an email and reply to it, which is a nightmare from an audit perspective. I think we need a dedicated view in the portal where managers can see all their pending approvals in one place — the full details of each request — and action them with an approve or reject and a comment. That should probably also show up as a widget on the homepage, but with a link through to a full-page view."),
    ("Tom Bakker", "That is crucial. Honestly that feature alone will get manager buy-in for the whole portal. Right now approval emails just get lost. I had a request sitting open for three weeks because someone's inbox was full."),
    ("Lisa Janssen", "So we're talking about a My Approvals view — filterable list, actionable inline, with approve and reject with a comment field?"),
    ("Mark de Vries", "And enough context per item — who submitted it, what they're asking for, when it came in. Not just a list of ticket numbers."),
    ("Tom Bakker", "Yes. And overdue items should stand out visually. If something has been waiting more than three working days, I want it flagged somehow."),
    ("Lisa Janssen", "So: My Approvals page with a filterable list of pending approvals including request details, inline approve and reject actions with a mandatory comment, and visual highlighting for items overdue by more than three working days. That's clean."),
    ("Sarah Chen", "Perfect. I think that's our four core areas. Let me read them back quickly."),
    ("Sarah Chen", "One — personalized homepage: announcements widget, open requests summary, role-based quick links. Two — service request catalog: category browsing, dynamic forms, file attachments, notes field. Three — knowledge base: searchable, categorized, ownership per category, helpfulness rating, suggested articles during request submission. Four — My Approvals dashboard for managers: pending approvals list with details, approve and reject with comment, overdue highlighting. Does that cover what we discussed?"),
    ("Tom Bakker", "Yes, that's it. Good summary. Feels like a solid v1 scope."),
    ("Priya Sharma", "Agreed from my side. Just making sure everyone remembers — the brand guidelines apply across all of it. Colors, fonts, button styles. I'll send the style guide to Mark this week."),
    ("Mark de Vries", "Thanks, looking forward to it."),
    ("Lisa Janssen", "Sarah, should we do a quick review session before the stories go into ServiceNow? I want to make sure the acceptance criteria are solid before the team starts picking things up."),
    ("Sarah Chen", "Good idea. I'll send an invite for next week. Tom, Priya — can you join for that review?"),
    ("Tom Bakker", "Should be fine, yes."),
    ("Priya Sharma", "I'll check my calendar but I'll make it work."),
    ("Sarah Chen", "Great. I'll get these captured as user stories by end of week. Thanks everyone — good session."),
    ("Mark de Vries", "Thanks all, bye!"),
    ("Tom Bakker", "Bye everyone."),
    ("Priya Sharma", "Bye, thanks!"),
    ("Lisa Janssen", "Bye!"),
]

PLAT4MATION_COLOR = RGBColor(0x1A, 0x73, 0xE8)   # blue for Plat4mation speakers
MERIDIAN_COLOR    = RGBColor(0x0F, 0x6B, 0x5E)   # teal for Meridian speakers

PLAT4MATION_SPEAKERS = {p[0] for p in PARTICIPANTS if p[2] == "Plat4mation"}


def add_speaker_line(doc, speaker, text):
    para = doc.add_paragraph()
    run_name = para.add_run(f"{speaker}: ")
    run_name.bold = True
    color = PLAT4MATION_COLOR if speaker in PLAT4MATION_SPEAKERS else MERIDIAN_COLOR
    run_name.font.color.rgb = color
    run_name.font.size = Pt(10.5)
    run_text = para.add_run(text)
    run_text.font.size = Pt(10.5)
    para.paragraph_format.space_after = Pt(4)
    return para


def generate_docx(output_path):
    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # --- Title ---
    title = doc.add_heading(MEETING_META["title"], level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.runs[0].font.size = Pt(16)

    # --- Metadata block ---
    meta = doc.add_paragraph()
    meta.add_run(f"Date: ").bold = True
    meta.add_run(f"{MEETING_META['date']}    ")
    meta.add_run(f"Duration: ").bold = True
    meta.add_run(f"{MEETING_META['duration']}    ")
    meta.add_run(f"Platform: ").bold = True
    meta.add_run(MEETING_META["platform"])
    meta.paragraph_format.space_after = Pt(8)

    # --- Participants ---
    doc.add_heading("Participants", level=2)
    for name, role, org in PARTICIPANTS:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{name}")
        run.bold = True
        p.add_run(f" — {role}, {org}")
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()  # spacer

    # --- Transcript ---
    doc.add_heading("Transcript", level=2)

    for speaker, text in TRANSCRIPT:
        add_speaker_line(doc, speaker, text)

    # --- Save ---
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"DOCX created: {output_path}")


if __name__ == "__main__":
    generate_docx(OUTPUT_PATH)
